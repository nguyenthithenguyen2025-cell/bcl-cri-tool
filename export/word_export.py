# -*- coding: utf-8 -*-
"""Xuất báo cáo kỹ thuật BCL-CRI ra file Word (.docx)."""

from io import BytesIO
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from config.parameters import PARAMETERS, PARAM_BY_ID
from config.parameters import RISK_COLORS
from core.classifier import generate_technical_analysis

RISK_LEVEL_LABELS = {
    1: "Cấp 1 — Rủi ro thấp",
    2: "Cấp 2 — Rủi ro trung bình",
    3: "Cấp 3 — Rủi ro cao",
    4: "Cấp 4 — Rủi ro rất cao",
}

RISK_HEX = {
    1: RGBColor(0x2e, 0xcc, 0x71),
    2: RGBColor(0xf3, 0x9c, 0x12),
    3: RGBColor(0xe6, 0x7e, 0x22),
    4: RGBColor(0xe7, 0x4c, 0x3c),
}


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(0x1f, 0x77, 0xb4)
    return h


def _add_table_row(table, label, value, bold_label=True):
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]
    label_cell.text = label
    value_cell.text = str(value) if value is not None else "—"
    if bold_label:
        label_cell.paragraphs[0].runs[0].font.bold = True
    return row


def export_to_word(
    entry: dict,
    include_solution: bool = True,
    include_legal: bool = True,
) -> BytesIO:
    """
    Tạo báo cáo kỹ thuật Word cho một BCL.

    Returns:
        BytesIO chứa file .docx
    """
    doc = Document()

    # ── Cài đặt trang A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    info = entry.get("info", {})
    result = entry.get("result", {})
    scores = entry.get("scores", {})
    missing_notes = entry.get("missing_notes", {})
    risk = result.get("risk", {})
    solution = result.get("solution", {})
    analysis = generate_technical_analysis(entry)

    # ══════════════════════════════════════════
    # TIÊU ĐỀ
    # ══════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO ĐÁNH GIÁ RỦI RO VÀ GIẢI PHÁP ĐÓNG BÃI CHÔN LẤP")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1f, 0x77, 0xb4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"Chỉ số Rủi ro Tổng hợp (CRI) — Đề tài TNMT.2024.05.05"
    ).font.size = Pt(11)

    doc.add_paragraph(f"Ngày lập báo cáo: {date.today().strftime('%d/%m/%Y')}").runs[0].font.italic = True
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # PHẦN 1 — THÔNG TIN BCL
    # ══════════════════════════════════════════
    _heading(doc, "1. Thông tin bãi chôn lấp", level=1)

    table_info = doc.add_table(rows=0, cols=2)
    table_info.style = "Table Grid"
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_info.columns[0].width = Cm(7)
    table_info.columns[1].width = Cm(10)

    fields = [
        ("Tên bãi chôn lấp", info.get("ten_bcl")),
        ("Tỉnh/Thành phố", info.get("tinh")),
        ("Xã/Phường", info.get("xa")),
        ("Loại BCL", "Không hợp vệ sinh (BCL-KHVS)" if info.get("loai_bcl") == "KHVS" else "Hợp vệ sinh (BCL-HVS)"),
        ("Diện tích (ha)", info.get("dien_tich_ha")),
        ("Thể tích ước tính (m³)", info.get("the_tich_m3")),
        ("Chiều cao ước tính (m)", info.get("chieu_cao_m")),
        ("Năm bắt đầu hoạt động", info.get("nam_bat_dau")),
        ("Năm ngừng tiếp nhận", info.get("nam_ngung")),
        ("Ghi chú", info.get("ghi_chu")),
    ]
    for label, value in fields:
        _add_table_row(table_info, label, value)

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # PHẦN 2 — KẾT QUẢ CRI
    # ══════════════════════════════════════════
    if info.get("loai_bcl") == "KHVS":
        _heading(doc, "2. Kết quả tính toán Chỉ số Rủi ro Tổng hợp (CRI)", level=1)

        cri_val = result.get("CRI")
        lv = risk.get("level")

        # Tóm tắt kết quả
        p_result = doc.add_paragraph()
        run = p_result.add_run(
            f"CRI = {cri_val:.4f}    "
            f"{risk.get('label', '')}    "
            f"Giải pháp: {solution.get('short_name', '')}"
            if cri_val else "Không tính được CRI."
        )
        run.font.bold = True
        run.font.size = Pt(12)
        if lv and lv in RISK_HEX:
            run.font.color.rgb = RISK_HEX[lv]

        doc.add_paragraph()

        # Bảng kết quả nhóm H, P, R, CRI
        _heading(doc, "2.1. Chỉ số từng nhóm", level=2)
        t_cri = doc.add_table(rows=1, cols=4)
        t_cri.style = "Table Grid"
        for i, h_text in enumerate(["Nhóm", "Chỉ số", "Mô tả", "Trọng số nhóm"]):
            cell = t_cri.rows[0].cells[i]
            cell.text = h_text
            cell.paragraphs[0].runs[0].font.bold = True

        for group, val, desc, exp in [
            ("H — Nguồn nguy hại", result.get("H"), "Đặc tính nguồn ô nhiễm", "0,28"),
            ("P — Đường lan truyền", result.get("P"), "Khả năng di chuyển ô nhiễm", "0,40"),
            ("R — Đối tượng tiếp nhận", result.get("R"), "Mức độ nhạy cảm của đối tượng", "0,32"),
            ("CRI tổng hợp", result.get("CRI"), "H^0,28 × P^0,40 × R^0,32", "—"),
        ]:
            row = t_cri.add_row()
            row.cells[0].text = group
            row.cells[1].text = f"{val:.4f}" if val else "—"
            row.cells[2].text = desc
            row.cells[3].text = exp
            if group == "CRI tổng hợp" and lv and lv in RISK_HEX:
                for c in row.cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.bold = True

        doc.add_paragraph()

        # Bảng điểm 14 thông số
        _heading(doc, "2.2. Điểm 14 thông số CRI", level=2)

        t_scores = doc.add_table(rows=1, cols=5)
        t_scores.style = "Table Grid"
        for i, h_text in enumerate(["Mã", "Tên thông số", "Nhóm", "Điểm", "Lý do thiếu dữ liệu"]):
            t_scores.rows[0].cells[i].text = h_text
            t_scores.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

        for p in PARAMETERS:
            pid = p["id"]
            score = scores.get(pid)
            assumed = pid in result.get("assumed_max", [])
            val_str = str(score) if score is not None else f"1,00 (*)"

            row = t_scores.add_row()
            row.cells[0].text = pid
            row.cells[1].text = p["name"]
            row.cells[2].text = p["group"]
            row.cells[3].text = val_str
            row.cells[4].text = missing_notes.get(pid, "—") if assumed else "—"
            if assumed:
                for c in row.cells:
                    for para in c.paragraphs:
                        for r in para.runs:
                            r.font.italic = True

        doc.add_paragraph("(*) Thông số được gán điểm 1,00 do thiếu dữ liệu (nguyên tắc thận trọng).").runs[0].font.italic = True
        doc.add_paragraph()

        _heading(doc, "3. Phân tích chuyên môn", level=1)
        doc.add_paragraph(analysis.get("summary", ""))

        _heading(doc, "3.1. Nhận xét theo nhóm H/P/R", level=2)
        for item in analysis.get("group_comments", []):
            doc.add_paragraph(
                f"{item['group']} — {item['name']}: {item['comment']}",
                style="List Bullet",
            )
        if analysis.get("dominant_group"):
            doc.add_paragraph(analysis["dominant_group"]).runs[0].font.bold = True

        _heading(doc, "3.2. Thông số chi phối rủi ro", level=2)
        for item in analysis.get("top_risks", [])[:5]:
            doc.add_paragraph(
                f"{item['id']} — {item['name']}: điểm {item['score']}, "
                f"đóng góp {item['contribution']:.4f}",
                style="List Bullet",
            )

        if analysis.get("data_quality_notes"):
            _heading(doc, "3.3. Chất lượng dữ liệu", level=2)
            for note in analysis["data_quality_notes"]:
                doc.add_paragraph(note, style="List Bullet")

        _heading(doc, "3.4. Khuyến nghị kỹ thuật tiếp theo", level=2)
        for action in analysis.get("recommended_actions", []):
            doc.add_paragraph(action, style="List Bullet")

        doc.add_paragraph()

    # ══════════════════════════════════════════
    # PHẦN 3 — GIẢI PHÁP KHUYẾN NGHỊ
    # ══════════════════════════════════════════
    if include_solution and solution:
        _heading(doc, "4. Giải pháp đóng bãi khuyến nghị", level=1)
        doc.add_paragraph(solution.get("name", "")).runs[0].font.bold = True
        doc.add_paragraph(solution.get("description", ""))

        _heading(doc, "4.1. Hạng mục bắt buộc", level=2)
        for item in solution.get("mandatory_items", []):
            p = doc.add_paragraph(item, style="List Bullet")

        _heading(doc, "4.2. Hạng mục khuyến nghị thêm", level=2)
        for item in solution.get("optional_items", []):
            doc.add_paragraph(item, style="List Bullet")

        _heading(doc, "4.3. Thời gian quản lý sau đóng bãi", level=2)
        doc.add_paragraph(solution.get("monitoring_period", "—"))

        doc.add_paragraph()

    # ══════════════════════════════════════════
    # PHẦN 4 — CĂN CỨ PHÁP LÝ
    # ══════════════════════════════════════════
    if include_legal:
        _heading(doc, "5. Căn cứ pháp lý", level=1)
        legal_items = [
            "Điều 32, Thông tư 02/2022/TT-BTNMT ngày 10/01/2022 của Bộ Tài nguyên và Môi trường",
            "QCVN 96:2025/BNNMT — Quy chuẩn kỹ thuật quốc gia về bãi chôn lấp chất thải rắn",
            "TCVN 13766:2023 — Chất thải rắn — Bãi chôn lấp hợp vệ sinh — Yêu cầu thiết kế",
            "Đề tài TNMT.2024.05.05 — Trường Đại học Thủy Lợi (2026) — Phương pháp CRI",
        ]
        for item in legal_items:
            doc.add_paragraph(item, style="List Bullet")

        if solution:
            doc.add_paragraph()
            doc.add_paragraph(f"Căn cứ áp dụng cho giải pháp: {solution.get('legal_basis', '—')}")

    # ── Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Báo cáo được tạo tự động bởi BCL-CRI Decision Support Tool v1.0 — "
        f"{date.today().strftime('%d/%m/%Y')}. "
        "Kết quả mang tính hỗ trợ quyết định, không thay thế đánh giá kỹ thuật chuyên ngành."
    )
    footer_p.runs[0].font.italic = True
    footer_p.runs[0].font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
